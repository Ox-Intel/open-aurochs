import base64
import datetime
import json
import io
import logging
import requests
import ast
import random
import re
import subprocess
import time
import traceback
import pypdf

# import pdfplumber
from concurrent.futures import ThreadPoolExecutor
from rapidfuzz import fuzz
from functools import partial
from bs4 import BeautifulSoup

from docx import Document

# import fitz
# from pdfminer.high_level import extract_text
from pdf2image import convert_from_bytes
from pytesseract import image_to_string
from docx2python import docx2python

from langchain_core.prompts import PromptTemplate
from langchain_openai import OpenAI
from langchain.chains import LLMChain
from langchain_openai import ChatOpenAI
from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
from langchain_community.agent_toolkits.load_tools import load_tools
from langchain.agents import initialize_agent
from langchain.agents import AgentType
from django.conf import settings
from celery import current_app
from django.contrib.contenttypes.models import ContentType
from django.utils import timezone
from asgiref.sync import async_to_sync
from channels.layers import get_channel_layer
from metaphor_python import Metaphor


from frameworks.models import Criteria, Framework
from organizations.models import (
    Organization,
    Team,
    TeamMember,
    User,
    OrganizationRole,
    Tag,
    TaggedObject,
    GenericPermission,
)
from reports.models import Report, Scorecard, ScorecardScore
from sources.models import Source, SourceFeedback
from stacks.models import Stack
from collaboration.models import ObjectSubscription, Comment, InboxItem
from history.models import ObjectHistoryChange


metaphor = Metaphor(settings.METAPHOR_KEY)

GPT_MODEL = "gpt-5.5"
GPT_LARGE_CONTEXT_MODEL = "gpt-5.4-nano"


PAGE_CONTENT_CUTOFF = 10000

NEWS_DOMAINS = [
    "cnn.com",
    "nytimes.com",
    "theguardian.com",
    "washingtonpost.com",
    "reuters.com",
    "npr.org",
    "wsj.com",
    "bloomberg.com",
    "ft.com",
    "apnews.com",
    "latimes.com",
    "thetimes.co",
    "aljazeera.com",
    "politico.com",
    "abcnews.com",
    "nbcnews.com",
    "cbsnews.com",
    "usatoday.com",
    "thehill.com",
    "dw.com",
    "afp.com",
    "scmp.com",
    "canadianpress.com",
    "ajc.com",
    "timesofindia.indiatimes",
    "independent.co",
    "time.com",
    "smh.com",
    "france24.com",
    "news.sky",
    "news.yahoo",
    "msn.com",
    "vox.com",
    "buzzfeednews.com",
    "businessinsider.com",
    "cnbc.com",
    "telegraph.co",
    "dailymail.co",
    "thehindu.com",
    "japantimes.co",
    "theglobeandmail.com",
    "bbc.co",
    "euronews.com",
    "nydailynews.com",
    "ndtv.com",
    "theconversation.com",
    "asahi.com",
    "economist.com",
    "axios.com",
]


class Mock(object):
    pass


def chunk_string(s, chunk_size=14000):
    return [s[i : i + chunk_size] for i in range(0, len(s), chunk_size)]


def research_chunk(chunk_info):
    summarize_context_chat = ChatOpenAI(model=GPT_LARGE_CONTEXT_MODEL)
    chunk, chunk_num, num_chunks, question, user_channel_id, gpt_guid = chunk_info
    time.sleep(chunk_num / num_chunks)

    pct = round(100 * (chunk_num - 1) / num_chunks)
    if pct == 100:
        pct = 98
    data = {
        "message": f"Compiling and evaluating research... ({chunk_num}/{num_chunks})",
        "inset": True,
        "complete": False,
        "event_type": "analyze-status-update",
    }
    return_oxgpt.delay(data, user_channel_id, gpt_guid)
    context_messages = [
        SystemMessage(content="You are a helpful assistant."),
        HumanMessage(
            content=f"Below are a number of accurate facts that provide context to a question about {question}. "
        ),
        HumanMessage(content=chunk),
        HumanMessage(
            content="That's all the facts. Now, summarize all of the above facts into a dense summary that is not more than 2000 words.  Focus on retaining factual details, and producing an accurate summary."
        ),
    ]
    c = summarize_context_chat.invoke(context_messages).content + "\n"
    # print(c)

    # chunk_num += 1
    data = {
        "message": f"Compiling and evaluating research... ({chunk_num}/{num_chunks})",
        "inset": True,
        "complete": True,
        "event_type": "analyze-status-update",
    }
    return_oxgpt.delay(data, user_channel_id, gpt_guid)
    return c


def gpt_summarize(messages, question="", user_channel_id="", gpt_guid=""):
    # print("gpt_summarize")
    summary = ""

    full_context = ""
    for m in messages:
        full_context += m + "\n\n"

    # print(len(full_context))
    if len(full_context) > 30000:
        chunks = chunk_string(full_context, chunk_size=30000)
        chunk_num = 1
        num_chunks = len(chunks)
        # print(f"chunks: {len(chunks)}")
        chunk_groups = []

        for c in chunks:
            chunk_groups.append(
                [c, chunk_num, num_chunks, question, user_channel_id, gpt_guid]
            )
            chunk_num += 1

        with ThreadPoolExecutor(max_workers=3) as executor:
            results = list(executor.map(research_chunk, chunk_groups))

        summary = "\n".join(results)

        # print(summary)
        data = {
            "message": "Compiling and evaluating research... complete",
            "inset": True,
            "complete": True,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)
        return summary
    else:
        return full_context


from io import StringIO
from html.parser import HTMLParser


class MLStripper(HTMLParser):
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = StringIO()

    def handle_data(self, d):
        self.text.write(d)

    def get_data(self):
        return self.text.getvalue()


def strip_tags(html):
    if html:
        s = MLStripper()
        s.feed(html)
        text = s.get_data()
        p = re.compile(r"<.*?>")
        return p.sub("", text)
    return ""


@current_app.task
def update_affected_users(object_type, object_pk, event_user_id):
    from api.events import event_handlers

    logging.info(object_type)
    logging.info(object_pk)
    # channel_layer = get_channel_layer()
    # async_to_sync(channel_layer.group_send)('core-data', message)

    cts = ContentType.objects.all().values("id", "app_label", "model")
    ct_framework_id = None
    ct_report_id = None
    ct_source_id = None
    ct_stack_id = None
    ct_scorecard_id = None
    ct_comment_id = None
    for ct in cts:
        if ct["model"] == "framework":
            ct_framework_id = ct["id"]
        if ct["model"] == "report":
            ct_report_id = ct["id"]
        if ct["model"] == "source":
            ct_source_id = ct["id"]
        if ct["model"] == "stack":
            ct_stack_id = ct["id"]
        if ct["model"] == "scorecard":
            ct_scorecard_id = ct["id"]
        if ct["model"] == "comment":
            ct_comment_id = ct["id"]
        if ct["model"] == "user":
            ct_user_id = ct["id"]

    content_types = {
        "framework": ct_framework_id,
        "report": ct_report_id,
        "source": ct_source_id,
        "stack": ct_stack_id,
        "scorecard": ct_scorecard_id,
        "comment": ct_comment_id,
        "user": ct_user_id,
    }
    content_types_by_id = {}
    content_types_by_id[str(ct_framework_id)] = "framework"
    content_types_by_id[str(ct_report_id)] = "report"
    content_types_by_id[str(ct_source_id)] = "source"
    content_types_by_id[str(ct_stack_id)] = "stack"
    content_types_by_id[str(ct_scorecard_id)] = "scorecard"
    content_types_by_id[str(ct_comment_id)] = "comment"
    content_types_by_id[str(ct_user_id)] = "user"

    all_data = False
    obj = None
    if object_type == "stack":
        obj = Stack.objects.get(pk=object_pk)
    elif object_type == "source":
        obj = Source.objects.get(pk=object_pk)
    elif object_type == "report":
        obj = Report.objects.get(pk=object_pk)
    elif object_type == "framework":
        obj = Framework.objects.get(pk=object_pk)
    elif object_type == "scorecard":
        sc = Scorecard.objects.get(pk=object_pk)
        obj = sc.report
        object_type = "report"
        object_pk = obj.pk
        all_data = True
    elif object_type == "comment":
        c = Comment.objects.get(pk=object_pk)
        object_type = content_types_by_id[str(c.content_type_id)]
        object_pk = c.object_id
        obj = c.content_object
        all_data = True

    if object_type == "inboxitem":
        return

    if not obj:
        if not (
            object_type == "organization"
            or object_type == "team"
            or object_type == "user"
        ):
            raise Exception("Unknown object type: %s" % object_type)

    if not (
        object_type == "organization" or object_type == "team" or object_type == "user"
    ):
        # Get all users, teams, orgs that can access it.
        # logging.error("getting all users who know about this...")
        user_ids = list(
            GenericPermission.objects.filter(
                object_id=object_pk, content_type_id=content_types[object_type]
            )
            .exclude(user=None)
            .distinct()
            .values_list("user_id", flat=True)
        )
        team_ids = (
            GenericPermission.objects.filter(
                object_id=object_pk, content_type_id=content_types[object_type]
            )
            .exclude(team=None)
            .distinct()
            .values_list("team_id", flat=True)
        )
        org_ids = (
            GenericPermission.objects.filter(
                object_id=object_pk, content_type_id=content_types[object_type]
            )
            .exclude(organization=None)
            .distinct()
            .values_list("organization_id", flat=True)
        )

        # logging.error(user_ids)
        team_user_ids = list(
            TeamMember.objects.filter(team_id__in=team_ids).values_list(
                "user_id", flat=True
            )
        )
        # logging.error(team_user_ids)
        org_user_ids = list(
            OrganizationRole.objects.filter(organization_id__in=org_ids).values_list(
                "user_id", flat=True
            )
        )
        # logging.error(org_user_ids)

        all_user_ids = set(user_ids + team_user_ids + org_user_ids)
    else:
        if object_type == "organization":
            o = Organization.objects.get(pk=object_pk)
            all_user_ids = [u.id for u in o.members]

        elif object_type == "team":
            t = Team.objects.get(pk=object_pk)
            all_user_ids = [u.id for u in t.members]

        elif object_type == "user":
            all_user_ids = [
                object_pk,
            ]

    # if event_user_id in all_user_ids:
    #     all_user_ids.remove(event_user_id)
    # logging.error(all_user_ids)

    # Update all the users
    channel_layer = get_channel_layer()
    users = User.objects.filter(id__in=all_user_ids).all()
    for user in users:
        request = Mock()
        request.user = user
        # data["_server_timestamp"] = timezone.now().timestamp() * 1000
        # if data["event_type"] in event_handlers:
        data = {}
        if event_user_id == user.id:
            all_data = True
        ret = event_handlers["get_my_user"]().handle(
            request,
            data,
            offline=True,
            all_data=all_data,
        )
        # # logging.error(ret)
        # logging.error(f"u-{user.ox_id}")
        # logging.error(ret)

        async_to_sync(channel_layer.group_send)(
            f"u-{user.ox_id}",
            {"type": "data_update", "data": ret},
        )
    # print(ret)
    return True


# @periodic_task(run_every=datetime.timedelta(minutes=30), expires=600)
# def send_pre_day_emails():
# pass


def split_text_into_chunks(text, chunk_size=6000):
    return [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]


def strip_non_numeric(s):
    return "".join([i for i in s if i.isdigit()])


def fetch_page_text(url):
    url = url.strip()
    # print(url)
    if url:
        # for unwanted in soup.find_all(['script', 'style', 'img', 'video', 'audio', 'aside', 'figure', 'footer', 'header', 'nav']):
        #     unwanted.extract()  # remove these tags and their contents
        #     url_contents[url] = soup.get_text()
        # /opt/homebrew/bin/wget -q -E -H -k -K -p -N --exclude-domains googletagmanager.com,cloudflareinsights.com,platform.linkedin.com --adjust-extension -e robots=off $URL
        # /opt/homebrew/bin/links -g -width 120 -codepage utf-8 -force-html -dump "index.html"
        # command = f"elinks -dump-charset utf-8 -force-html -dump {url}"
        command = f"links -g -width 120 -codepage utf-8 -force-html -dump {url}"
        process = subprocess.Popen(command, stdout=subprocess.PIPE, shell=True)
        output, error = process.communicate()

        if error:
            print(f"An error occurred: {error}")
        else:
            url_content = output.decode()  # decode bytes to string

        # Get the title
        # command = f" curl -s {url} | grep -o '<title>[^<]*</title>' | sed 's/<title>\\(.*\\)<\\/title>/\1/g'"
        # process = subprocess.Popen(command, stdout=subprocess.PIPE, shell=True)
        # output, error = process.communicate()
        resp = requests.get(url)
        soup = BeautifulSoup(resp.content, "html.parser")
        title = soup.title.string if soup.title else url

        # title = output.decode()
        return (url, url_content, title)
    return None


def score_single_criteria_url(pairs):
    url, url_content, c, question, context, title = pairs
    # print(c, url)
    if c:
        name = c["name"]
        content = "You are a helpful assistant that evaluates how well a text source supports a given question."

        question_context = ""
        if question:
            question_context = f", in the context of {question}"

        source_context = f"The source below has page summaries from {context}"
        chat = ChatOpenAI(model=GPT_MODEL)
        messages = [
            SystemMessage(content=content),
            HumanMessage(
                content=f"Evaluate how strongly the following source scores for {name} {question_context}. {source_context}. Please ignore any navigation, headers, or footer links, as well as links to other pages.  In your rating, please do not evaluate the source itself, but only how the source describes {title} regarding {name} Please return a single number between 0 and 10 (inclusive of 0 and 10), where where 0 is a bad score and 10 is a very good score, followed by a |, followed by a brief explanation of why that score was given. Do not include anything besides the number score before the | character.  For example, \n8 | This is an explanation.\n\n Below is the source to evaluate."  # noqa
            ),
        ]
        chunks = split_text_into_chunks(url_content, chunk_size=4000)
        for i, chunk in enumerate(chunks):
            messages.append(HumanMessage(content=chunk))

        # print(len(url_content))
        # print(len(description))
        # print(len(messages))
        # print(url_contents[url])
        # raise Exception("stop")

        # print(messages)
        # chain = LLMChain(llm=llm, prompt=prompt)
        # scoring_response = chain.run({"name":name, "description":description, "url_content":url_content, "question":question,})

        scoring_response = chat.invoke(messages).content
        return c, url, scoring_response
    return None


def score_single_criteria_topic(pairs):
    topic, c, question, context, title = pairs
    # print(c, url)
    if c:
        name = c["name"]
        # weight = c["weight"]
        description = c["description"]
        if True:
            #         try:
            #             # Need SERP and agency.
            #             assert True is False
            #             llm = OpenAI(temperature=0.3, model=GPT_MODEL)
            #             tools = load_tools(["serpapi", "llm-math", "requests", "wikipedia"], llm=llm)
            #             # tools = load_tools(["serpapi", "python_repl"], llm=llm)
            #             agent = initialize_agent(
            #                 tools, llm, agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION, verbose=True
            #             )

            #             # scoring_response = agent.run(f"Using both the internet and your general knowledge, evaluate how strongly {topic} scores for {name}, in helping the user {question}. {context}.  If you run into trouble searching or requesting a link, please ignore it and continue. In using wikipedia, make absolutely sure that the subject matches before including it.  Please return a single number between 1 and 10, where where 1 is very weakly supported and 10 is very strongly supported, followed by a |, followed by a brief explanation of why that score was given. Cite all URLs or sources used to make the assessment. Do not include anything besides the number score before the | character.  For example, \n8 | This is an explanation." )
            #             scoring_response = agent.run(
            #                 f"""Please use all of: serp and the internet, wikipedia (only if an exact match is found) and your general knowledge.   If you run into trouble searching or requesting a link, please ignore it and continue. In using wikipedia, make absolutely sure that the subject matches before including it.  Evaluate how strongly {topic} scores for {name}, in helping the user {question}. {context}.  Please return a single number between 1 and 10, where where 1 is very weakly supported and 10 is very strongly supported, followed by a |, followed by a brief explanation of why that score was given. Cite any URLs used to make the assessment. Do not include anything besides the number score before the | character.  For example,
            # 8 | This is an explanation.

            # Assume that {topic} is a real thing, not a speculative idea or thought experiment.  Please use only your knowledge of the world to answer the above prompt.  If you don't have any information about {topic} and how it scores for {name}, please say that you don't know. In these cases, please do not include apologies, or explanations that you are an AI assistant - simply state that there is not enough available information to answer the question."""
            #             )
            #             # print(scoring_response)
            #         except:
            # Fall back to Zero-shot chat.
            content = f"You are a helpful, AI assistant with expert knowledge about the {question}, being asked to evaluate and score how well a subject scores for a given criteria."

            # Test to see what it's thinking behind the scenes.
            # chat = ChatOpenAI(temperature=0.1, model=GPT_MODEL)
            # messages = [
            #     SystemMessage(content=content),
            #     HumanMessage(
            #         content=f"Assume that {topic} is a real thing, not a speculative idea or thought experiment.  Please use your knoweledge of the world to answer the following question: What is {topic}?   If you don't know, please say that you don't know, but do not include apologies, or an explanation that you are an AI assistant."
            #     ),
            # ]
            # print(chat.invoke(messages).content)

            chat = ChatOpenAI(model=GPT_MODEL)
            prompt = f"""Evaluate how strongly {topic} scores for the criteria {name} (e.g {description}), in a framework focused on {question}. {context}.  Focus only on {name}, not related criteria or ideas, nor how important {name} is to deciding {question}.  Please return a list in the format below with the following keys: score, comment, skipped. Be accurate with your scores, and don't be afraid to use the full numerical range.  If you don't have enough information to answer, please return True for skipped.   For score, please return a single number between 0 and 10, where where 0 is not supported and 10 is completely supported.

Assume that {topic} is a real thing, not a speculative idea or thought experiment.  Please use only your knowledge of the world and the context above to answer the prompt.  If you don't have any information about {topic} and how it scores for {name}, please say that you don't know, return skipped as True, and return a score of -1.

When replying, please use the following format, separating each field by a line break, and starting with a |.
|score: 5
|comment: a string containing a brief explanation of why that score was given.
|skipped: False
"""

            # Both comment and confidence_comment should be no longer than one short sentence.

            # Please return a single number between 0 and 10, where where 1 is not supported and 10 is completely supported.   and return the number -1, using the same formatting.   When returning the output, follow the number with a |, followed by a brief explanation of why that score was given. Do not include anything besides the number score before the | character.  For example:
            # 8 | This is an explanation.
            # Assume that {topic} is a real thing, not a speculative idea or thought experiment.  Please use only your knowledge of the world to answer the above prompt.  If you don't have any information about {topic} and how it scores for {name}, please say that you don't know. In these cases, please do not include apologies, or explanations that you are an AI assistant - simply state that there is not enough available information to answer the question and return a score of -1."""

            messages = [
                SystemMessage(content=content),
                HumanMessage(content=prompt),
            ]
            # print(prompt)
            scoring_response1 = chat.invoke(messages).content

            parsed = (
                scoring_response1.replace(": false,", ": False,")
                .replace(": true,", ": True,")
                .split("|")[1:]
            )
            score = int(strip_non_numeric(parsed[0].split(":")[-1]))
            comment = parsed[1].split(":")[-1]
            try:
                skipped = ast.literal_eval(parsed[2].split(":")[-1].strip())
            except:
                traceback.print_exc()

            parsed_response = {
                "score": score,
                "comment": comment,
                "skipped": skipped,
            }
            # messages = [
            #     SystemMessage(content=content),
            #     HumanMessage(
            #         content=prompt
            #     ),
            # ]
            # # print(prompt)
            # scoring_response2 = chat.invoke(messages).content

            # print(scoring_response1)
        #             chat = ChatOpenAI(temperature=0.1, model=GPT_MODEL)
        #             prompt = f"""Reformat and combine the following responses into one valid python dictionary with the following keys: score, comment, confidence, confidence_comment, and skipped, based on the following response below inside the | characters.  Please don't return any commentary or code blocks, just the dictionary.  If score is missing or null, please set it to -1, and skipped to True, and make sure that a comment is included.  If the comment includes explanatory text about why a score was skipped, please remove it.
        # |{scoring_response1}|
        # """

        # Please return a single number between 0 and 10, where where 1 is not supported and 10 is completely supported.   and return the number -1, using the same formatting.   When returning the output, follow the number with a |, followed by a brief explanation of why that score was given. Do not include anything besides the number score before the | character.  For example:
        # 8 | This is an explanation.
        # Assume that {topic} is a real thing, not a speculative idea or thought experiment.  Please use only your knowledge of the world to answer the above prompt.  If you don't have any information about {topic} and how it scores for {name}, please say that you don't know. In these cases, please do not include apologies, or explanations that you are an AI assistant - simply state that there is not enough available information to answer the question and return a score of -1."""

        # messages = [
        #     SystemMessage(content=content),
        #     HumanMessage(content=prompt),
        # ]
        # engineered_response = (
        #     chat(messages)
        #     .content.replace(": false,", ": False,")
        #     .replace(": true,", ": True,")
        # )
        # # print(engineered_response)

        # trimmed_response = (
        #     engineered_response[
        #         engineered_response.find("{") : engineered_response.rfind("}") + 1
        #     ]
        #     .replace('"', '"')
        #     .replace("|", "")
        # )
        # # scoring_response = f"[{scoring_response}]"
        # try:
        #     trimmed_response = ast.literal_eval(trimmed_response)
        # except:
        #     from traceback import print_exc

        #     print("failed to parse")
        #     # print(scoring_response)
        #     print(trimmed_response)
        #     print_exc()
        #     raise Exception("failed")

        # print("parsed_response")
        # print(parsed_response)
        # print("trimmed_response")
        # print(trimmed_response)
        # assert parsed_response != trimmed_response

        # print("trimmed_response:")
        # print(trimmed_response)

        #             prompt = f"""Evaluate how much confididence a hypothetical human expert who has the same real-world knowledge that you do would have in a score of {parsed_response['score']} for {topic}, when scoring for the criteria {name} (e.g {description}), in a framework focused on {question}. {context}. Please return a list in the format below with the following keys: confidence, confidence_comment. For score the range is from 0 and 10, where where 0 is not supported and 10 is completely supported.

        # For confidence, please return a single number between 0 and 10 that indicates how confident you are that you have scored this criteria accurately, based on how much knowledge you have about {topic} and {name} (be conservative) as an AI relative to how much knowledge an expert would have, where where 1 is not at all confident and 10 is completely confident.  Use comment to explain why you chose that score, and confidence_comment to explain why you chose that confidence.  confidence_comment should be no longer than one short sentence.

        # Assume that {topic} is a real thing, not a speculative idea or thought experiment.  Please use only your knowledge of the world to answer the above prompt.

        # When replying, please use the following format, separating each field by a line break, and starting with a |.
        # |confidence: 3
        # |confidence_comment: a short sentence containing a brief explanation of why confidence was given.
        # """
        #             messages = [
        #                 SystemMessage(content=content),
        #                 HumanMessage(content=prompt),
        #             ]
        #             confidence_response = chat.invoke(messages).content
        #             parsed = (
        #                 confidence_response.replace(": false,", ": False,")
        #                 .replace(": true,", ": True,")
        #                 .split("|")[1:]
        #             )
        #             confidence = parsed[0].split(":")[1]
        #             confidence_comment = parsed[1].split(":")[1]
        #             parsed_response["confidence"] = int(strip_non_numeric(confidence))
        #             parsed_response["confidence_comment"] = confidence_comment

        #             prompt = f"""Please return a number between 0 and 10 representing how strong the following statement is, where 0 is not at all confident, and 10 is completely confident.  Please return only that number, with no explanatory text or punctuation.

        # {confidence_comment}
        # """
        #             messages = [
        #                 SystemMessage(content=content),
        #                 HumanMessage(content=prompt),
        #             ]
        #             confidence_response = strip_non_numeric(chat.invoke(messages).content)
        #             parsed_response["confidence"] = int(confidence_response)

        return topic, c, parsed_response
    return None


def score_single_criteria_all_subjects(pairs):
    # print(f"calling score_single_criteria_all_subjects:")
    # Introduce entropy to work around rate limits.
    # time.sleep(random.randint(0, 100) / 30)

    subjects, c, question, full_context = pairs
    # source_general_knowledge = full_context["source_general_knowledge"]
    # source_search = full_context["source_search"]
    # source_news = full_context["source_news"]
    # source_custom = full_context["source_custom"]
    # source_urls = full_context["source_urls"]
    context_messages = full_context["context_messages"]
    context = full_context["scoring_context"]
    user_channel_id = full_context["user_channel_id"]
    gpt_guid = full_context["gpt_guid"]
    subjects_string = "\n".join(subjects)

    messages = []
    # print(len(context_messages))
    for m in context_messages:
        if m:
            messages.append(SystemMessage(content=m))

    if c:
        name = c["name"]
        data = {
            "message": f"Scoring {name}...",
            "inset": True,
            "complete": False,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)

        # weight = c["weight"]
        description = c["description"]
        subjects_string = "\n".join(subjects)

        chat = ChatOpenAI(model=GPT_LARGE_CONTEXT_MODEL)
        prompt = f"""Ok.  So given that context, write a summary of how strongly the following list of subjects would score for the criteria {name} (e.g {description}), in a framework focused on "{question}". {context}.  Focus only on {name}, not related criteria or ideas, nor how important {name} is to "{question}".  Please compare the subjects against each other to achieve an accurate scoring for all. Do not use the order of subjects to weight your judgement.

Subjects: {subjects_string}

Please return a list that includes at least one paragraph for each subject, describing how strongly they score for {name}.  Make sure that you write a paragraph for every single subject.

Structure the list like this:
- Subject 1: Subject 1 is .....
- Subject 2: Subject 2 is ....."""

        messages.append(
            HumanMessage(content=prompt),
        )

        # print(prompt)
        summary_response = chat.invoke(messages).content

        # summary_response_truncated = chat.invoke(messages).content
        # If we're getting cutoffs.
        # messages.append(SystemMessage(content=f"AI Response: {summary_response_truncated}"))
        # print(summary_response_truncated)
        # messages.append(HumanMessage(content="Please double-check the response above. Were all subjects included - {subjects_string}? Make sure it is complete."))
        # if False:
        #     print(summary_response_truncated)
        # summary_response = chat.invoke(messages).content
        # print("summary_response")
        # print(summary_response)

        chat = ChatOpenAI(model=GPT_MODEL)
        prompt = f"""Here is some context for the following question:
Evaluation of {name} for each subject:
{summary_response}

Using the above context only, evaluate how strongly the following list of subjects score for the criteria {name} (e.g {description}), in a framework focused on "{question}". Focus only on {name}, not related criteria or ideas, nor how important {name} is to "{question}".  Please compare the subjects against each other to achieve an accurate scoring for all. Do not use the order of subjects to weight your judgement.

Please return a list in the format below with the following keys: score, comment, skipped, speculation. Be accurate with your scores, and don't be afraid to use the full numerical range.  If you don't have enough information to answer, please return True for skipped.   For score, please return a single number between 0 and 10, where where 0 is not supported and 10 is completely supported.  For speculation, please return a score of how speculative the sentence is, where 0 is highly speculative, and 10 is not at all speculative.

Assume that each subject is a real thing, not a speculative idea or thought experiment.  Do not make assumptions. Please use ONLY your knowledge of the world and the context given to answer the prompt, and do not make any implications, deductions, or assumptions from the subject names.  If it is unclear what the subject is or how it should be scored, return a low score between 0 and 4. If you don't have any real information about the subject and how it scores for {name}, please say that you don't know, return skipped as True, and return a score of -1.

Subjects:
{subjects_string}

When replying, please use the following format, separating each field by a line break, and starting with a |.  Put the subject between the == marks. Use the exact subject spelling from the list above.
=={subjects[0]}==
|score: 5
|comment: a string containing a brief explanation of why that score was given for {subjects[0]}.
|speculation: 2
|skipped: False

==etc==
|score: -1
|comment: a string containing a brief explanation of why that score was given for Subject 2.
|speculation: 8
|skipped: True
"""
        messages = []
        messages.append(
            HumanMessage(content=prompt),
        )

        # print(prompt)
        scoring_response = chat.invoke(messages).content

        if False:
            print(scoring_response)
        finding_subject = True
        scores = {}
        subject = ""
        orphaned_subjects = []
        while "==" in scoring_response:
            chunk = scoring_response[: scoring_response.find("==")].strip()
            if not chunk:
                scoring_response = scoring_response[scoring_response.find("==") + 2 :]
                chunk = scoring_response[: scoring_response.find("==")].strip()

            # print(f"finding: {finding_subject} chunk: |{chunk}|")
            if finding_subject:
                subject = chunk
                if subject not in subjects:
                    orphaned_subjects.append(subject)
                scores[subject] = {}
                finding_subject = False
            else:
                categories = chunk.split("|")
                for cat in categories:
                    cat = cat.strip()
                    # print(f"{subject} c: {c}")
                    if cat:
                        key = cat[: cat.find(":")].strip()
                        val = cat[cat.find(":") + 1 :].strip()
                        # print(f"{key} = {val}")
                        if key == "score" or key == "speculation":
                            scores[subject][key] = int(
                                strip_non_numeric(val.replace("\n", ""))
                            )
                        elif key == "skipped":
                            try:
                                scores[subject][key] = ast.literal_eval(
                                    val.replace("\n", "")
                                )
                            except:
                                traceback.print_exc()
                        else:
                            scores[subject][key] = val
                finding_subject = True

            scoring_response = scoring_response[scoring_response.find("==") + 2 :]

        chunk = scoring_response.strip()
        if finding_subject:
            subject = chunk
            if subject not in subjects:
                orphaned_subjects.append(subject)
            scores[subject] = {}
            finding_subject = False
        else:
            categories = chunk.split("|")
            for cat in categories:
                cat = cat.strip()
                # print(f"{subject} cat: {cat}")
                if cat:
                    key = cat[: cat.find(":")].strip()
                    val = cat[cat.find(":") + 1 :].strip()
                    # print(f"{key} = {val}")
                    if key == "score" or key == "speculation" and val != "":
                        scores[subject][key] = int(
                            strip_non_numeric(val.replace("\n", ""))
                        )
                    elif key == "skipped":
                        try:
                            scores[subject][key] = ast.literal_eval(
                                val.replace("\n", "")
                            )
                        except:
                            traceback.print_exc()
                    else:
                        scores[subject][key] = val

        # Handle cases where our LLM returns spelling corrections or the like.
        for os in orphaned_subjects:
            best_ratio = 0
            best_ratio_subject = None
            for s in subjects:
                ratio = fuzz.ratio(os, s)
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_ratio_subject = s
            # print(f"|{os}|, |{best_ratio_subject}|, {best_ratio}")
            if best_ratio_subject and best_ratio > 80:
                for k, v in scores[os].items():
                    if best_ratio_subject not in scores:
                        scores[best_ratio_subject] = {}
                    scores[best_ratio_subject][k] = v
                del scores[os]

        # print(scores)

        # Now, I want to add speculation scores.
        subject_strings = ""
        count = 1
        for subject, score in scores.items():
            if "comment" in score:
                subject_strings += f"{subject}: {score['comment']}\n"
                count += 1
            else:
                pass
                # print(score)
                # raise Exception

        #         prompt = f"""For each of the following sentences, return a score of how confident the sentence is about {name} ({description}) subject, where 0 is not at all confident, and 10 is extremely confident that it has real-world knowledge.

        # {subject_strings}

        # Please return the the sentence number, followed by a |, followed by the score, separated by line breaks.  For example:

        # 1. 2
        # 2. 8
        # 3. 3
        # etc.
        # """

        #         messages = [
        #             HumanMessage(content=prompt),
        #         ]
        #         # print(prompt)
        #         scoring_response = chat.invoke(messages).content
        #         print(scoring_response)
        #         confidence_scores = scoring_response.split('\n')
        #         count = 0
        #         for subject, score in scores.items():
        #             score["confidence"] = confidence_scores[count].split(". ")[1]
        #             count += 1

        chat = ChatOpenAI(model=GPT_MODEL)

        prompt = f"""Below is a list of topics and descriptive sentences.
For each topic, evaluate the descriptive sentence, based on your real-world knowledge about {name} in the context of {question} and the sentence content, and return a score of how speculative that sentence is, where 0 is entirely based on real facts, and 10 is highly speculative and fabricated without evidence.

When replying, output a valid python dictionary, where the key is the topic, and the value is the score.

For example:
topic one: this is a factually-based sentence.
second topic: this sentence might be about some useful things, but we've just made it up.

Should return:
{{
    "topic one": 2,
    "second topic": 8,
}}

Here are the topics and sentences to evaluate:
{subject_strings}

"""
        # print(prompt)
        messages = [
            HumanMessage(content=prompt),
        ]
        # print(prompt)
        scoring_response = chat.invoke(messages).content
        # print(scoring_response)
        # print(prompt)
        # scoring_response = scoring_response[scoring_response.find("[") : scoring_response.rfind("]") + 1]
        try:
            parsed_scores = ast.literal_eval(scoring_response)
        except:
            traceback.print_exc()
        # print(parsed_scores)
        # speculation_scores = scoring_response.split('\n')
        count = 0
        # print(f"scores: {len(scores)} spec scores: {len(parsed_scores)}")

        orphaned_scores = []
        for subject, score in scores.items():
            # print(count)
            # print(speculation_scores[count])
            # score["separate_speculation"] = int(strip_non_numeric(speculation_scores[count].split(":")[0]))
            found = False
            for k, s in parsed_scores.items():
                if k.lower() == subject.lower():
                    # print(f"{name}: {subject}-{s}")
                    score["separate_speculation"] = s
                    found = True
                    break

            if not found:
                orphaned_scores.append((k, s))

        for k, speculation_score in orphaned_scores:
            if speculation_score or speculation_score == 0:
                best_ratio = 0
                best_ratio_subject = None
                for s in subjects:
                    ratio = fuzz.ratio(k, s)
                    if ratio > best_ratio:
                        best_ratio = ratio
                        best_ratio_subject = s
                # print(f"|{os}|, |{best_ratio_subject}|, {best_ratio}")
                if best_ratio_subject and best_ratio > 80:
                    if best_ratio_subject not in scores:
                        scores[best_ratio_subject] = {}
                    scores[best_ratio_subject][
                        "separate_speculation"
                    ] = speculation_score

            # count += 1

        # print(scores)
        data = {
            "message": f"Scoring {name}...",
            "complete": True,
            "inset": True,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)

        return c, scores
    return None


def score_single_criteria_single_subjects(pairs):
    # print(f"calling score_single_criteria_all_subjects: {pairs}")
    # Introduce entropy to work around rate limits.
    # time.sleep(random.randint(0, 100) / 30)

    subjects, c, question, full_context = pairs
    # source_general_knowledge = full_context["source_general_knowledge"]
    # source_search = full_context["source_search"]
    # source_news = full_context["source_news"]
    # source_custom = full_context["source_custom"]
    # source_urls = full_context["source_urls"]
    context_messages = full_context["context_messages"]
    context = full_context["scoring_context"]
    user_channel_id = full_context["user_channel_id"]
    gpt_guid = full_context["gpt_guid"]
    subjects_string = "\n".join(subjects)

    messages = []
    for m in context_messages:
        if m:
            messages.append(SystemMessage(content=m))

    # print(c, url)
    if c:
        name = c["name"]
        data = {
            "message": f"Scoring {subjects} for {name}...",
            "complete": False,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)
        # weight = c["weight"]
        description = c["description"]
        subjects_string = "\n".join(subjects)

        chat = ChatOpenAI(model=GPT_LARGE_CONTEXT_MODEL)

        prompt = f"""Evaluate how strongly the following list of subjects score for the criteria {name} (e.g {description}), in a framework focused on "{question}". {context}.  Focus only on {name}, not related criteria or ideas, nor how important {name} is to "{question}".  Please compare the subjects against each other to achieve an accurate scoring for all. Do not use the order of subjects to weight your judgement.  Please return a list in the format below with the following keys: score, comment, skipped, speculation. Be accurate with your scores, and don't be afraid to use the full numerical range.  If you don't have enough information to answer, please return True for skipped.   For score, please return a single number between 0 and 10, where where 0 is not supported and 10 is completely supported.  For speculation, please return a score of how speculative the sentence is, where 0 is highly speculative, and 10 is not at all speculative.

Assume that each subject is a real thing, not a speculative idea or thought experiment.  Do not make assumptions. Please use ONLY your knowledge of the world and the context given to answer the prompt, and do not make any implications, deductions, or assumptions from the subject names.  If it is unclear what the subject is or how it should be scored, return a low score between 0 and 4. If you don't have any real information about the subject and how it scores for {name}, please say that you don't know, return skipped as True, and return a score of -1.

Subjects:
{subjects_string}

When replying, please use the following format, separating each field by a line break, and starting with a |.  Put the subject between the == marks. Use the exact subject spelling from the list above.
=={subjects[0]}==
|score: 5
|comment: a string containing a brief explanation of why that score was given for {subjects[0]}.
|speculation: 2
|skipped: False

==etc==
|score: -1
|comment: a string containing a brief explanation of why that score was given for Subject 2.
|speculation: 8
|skipped: True
"""

        messages.append(
            HumanMessage(content=prompt),
        )

        # print(prompt)
        scoring_response = chat.invoke(messages).content
        name = c["name"]

        # print(scoring_response)
        finding_subject = True
        scores = {}
        subject = ""
        orphaned_subjects = []
        while "==" in scoring_response:
            chunk = scoring_response[: scoring_response.find("==")].strip()
            if not chunk:
                scoring_response = scoring_response[scoring_response.find("==") + 2 :]
                chunk = scoring_response[: scoring_response.find("==")].strip()

            # print(f"finding: {finding_subject} chunk: |{chunk}|")
            if finding_subject:
                subject = chunk
                if subject not in subjects:
                    orphaned_subjects.append(subject)
                scores[subject] = {}
                finding_subject = False
            else:
                categories = chunk.split("|")
                for cat in categories:
                    cat = cat.strip()
                    # print(f"{subject} c: {c}")
                    if cat:
                        key = cat[: cat.find(":")].strip()
                        val = cat[cat.find(":") + 1 :].strip()
                        # print(f"{key} = {val}")
                        if key == "score" or key == "speculation":
                            scores[subject][key] = int(
                                strip_non_numeric(val.replace("\n", ""))
                            )
                        elif key == "skipped":
                            try:
                                scores[subject][key] = ast.literal_eval(
                                    val.replace("\n", "")
                                )
                            except:
                                traceback.print_exc()
                        else:
                            scores[subject][key] = val
                finding_subject = True

            scoring_response = scoring_response[scoring_response.find("==") + 2 :]

        chunk = scoring_response.strip()
        if finding_subject:
            subject = chunk
            if subject not in subjects:
                orphaned_subjects.append(subject)
            scores[subject] = {}
            finding_subject = False
        else:
            categories = chunk.split("|")
            for cat in categories:
                cat = cat.strip()
                # print(f"{subject} cat: {cat}")
                if cat:
                    key = cat[: cat.find(":")].strip()
                    val = cat[cat.find(":") + 1 :].strip()
                    # print(f"{key} = {val}")
                    if key == "score" or key == "speculation" and val != "":
                        scores[subject][key] = int(
                            strip_non_numeric(val.replace("\n", ""))
                        )
                    elif key == "skipped":
                        try:
                            scores[subject][key] = ast.literal_eval(
                                val.replace("\n", "")
                            )
                        except:
                            traceback.print_exc()
                    else:
                        scores[subject][key] = val

        # Handle cases where our LLM returns spelling corrections or the like.
        for os in orphaned_subjects:
            best_ratio = 0
            best_ratio_subject = None
            for s in subjects:
                ratio = fuzz.ratio(os, s)
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_ratio_subject = s
            # print(f"|{os}|, |{best_ratio_subject}|, {best_ratio}")
            if best_ratio_subject and best_ratio > 80:
                for k, v in scores[os].items():
                    if best_ratio_subject not in scores:
                        scores[best_ratio_subject] = {}
                    scores[best_ratio_subject][k] = v
                del scores[os]

        # print(scores)

        # Let's score these based on the comments.

        # Now, I want to add speculation scores.
        subject_strings = ""
        count = 1
        for subject, score in scores.items():
            if "comment" in score:
                subject_strings += f"{subject}: {score['comment']}\n"
                count += 1
            else:
                pass
                # print(score)
                # raise Exception

        #         prompt = f"""For each of the following sentences, return a score of how confident the sentence is about {name} ({description}) subject, where 0 is not at all confident, and 10 is extremely confident that it has real-world knowledge.

        # {subject_strings}

        # Please return the the sentence number, followed by a |, followed by the score, separated by line breaks.  For example:

        # 1. 2
        # 2. 8
        # 3. 3
        # etc.
        # """

        #         messages = [
        #             HumanMessage(content=prompt),
        #         ]
        #         # print(prompt)
        #         scoring_response = chat.invoke(messages).content
        #         print(scoring_response)
        #         confidence_scores = scoring_response.split('\n')
        #         count = 0
        #         for subject, score in scores.items():
        #             score["confidence"] = confidence_scores[count].split(". ")[1]
        #             count += 1

        chat = ChatOpenAI(model=GPT_MODEL)

        prompt = f"""Below is a list of topics and descriptive sentences.
For each topic, evaluate the descriptive sentence, based on your real-world knowledge about {name} in the context of {question} and the sentence content, and return a score of how speculative that sentence is, where 0 is entirely based on real facts, and 10 is highly speculative.

When replying, output a valid python dictionary, where the key is the topic, and the value is the score.

For example:
topic one: this is a factually-based sentence.
second topic: this sentence might be about some useful things, but we've just made it up.

Should return:
{{
    "topic one": 2,
    "second topic": 8,
}}

Here are the topics and sentences to evaluate:
{subject_strings}

"""
        # print(prompt)
        messages = [
            HumanMessage(content=prompt),
        ]
        # print(prompt)
        scoring_response = chat.invoke(messages).content
        # print(scoring_response)
        # print(prompt)
        # scoring_response = scoring_response[scoring_response.find("[") : scoring_response.rfind("]") + 1]
        try:
            parsed_scores = ast.literal_eval(scoring_response)
        except:
            traceback.print_exc()
        # print(parsed_scores)
        # speculation_scores = scoring_response.split('\n')
        count = 0
        # print(f"scores: {len(scores)} spec scores: {len(parsed_scores)}")

        orphaned_scores = []
        for subject, score in scores.items():
            # print(count)
            # print(speculation_scores[count])
            # score["separate_speculation"] = int(strip_non_numeric(speculation_scores[count].split(":")[0]))
            found = False
            for k, s in parsed_scores.items():
                if k.lower() == subject.lower():
                    # print(f"{name}: {subject}-{s}")
                    score["separate_speculation"] = s
                    found = True
                    break

            if not found:
                orphaned_scores.append((k, s))

        for k, speculation_score in orphaned_scores:
            if speculation_score or speculation_score == 0:
                best_ratio = 0
                best_ratio_subject = None
                for s in subjects:
                    ratio = fuzz.ratio(k, s)
                    if ratio > best_ratio:
                        best_ratio = ratio
                        best_ratio_subject = s
                # print(f"|{os}|, |{best_ratio_subject}|, {best_ratio}")
                if best_ratio_subject and best_ratio > 80:
                    if best_ratio_subject not in scores:
                        scores[best_ratio_subject] = {}
                    scores[best_ratio_subject][
                        "separate_speculation"
                    ] = speculation_score

            # count += 1

        # print(scores)
        data = {
            "message": f"Scoring {subjects} for {name}...",
            "complete": True,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)

        return c, scores
    return None


@current_app.task
def return_oxgpt(data, user_channel_id, gpt_guid):
    # print(f"return_oxgpt {user_channel_id}")
    # print(data)
    # Update all the users
    channel_layer = get_channel_layer()
    # user = User.objects.get(pk=user_id)
    data["sentAt"] = datetime.datetime.now().timestamp() * 1000
    data["gptInstanceGUID"] = gpt_guid

    async_to_sync(channel_layer.group_send)(
        f"u-{user_channel_id}",
        {"type": "oxgpt_event", "data": data},
    )
    # print(ret)
    return True


@current_app.task
def return_export(data, user_channel_id):
    # Update all the users
    channel_layer = get_channel_layer()
    # user = User.objects.get(pk=user_id)

    async_to_sync(channel_layer.group_send)(
        f"u-{user_channel_id}",
        {"type": "export_event", "data": data},
    )
    # print(ret)
    return True


@current_app.task
def oxgpt_generate_framework(data, user_channel_id):
    # print(f"oxgpt_generate_framework: {user_channel_id}")
    criteria = []
    question = data["topic"]
    framework_context = data["topic_context"]
    gpt_guid = data["gptInstanceGUID"]

    # for s in subject_list:
    #     if s.strip():
    #         if "://" in s:
    #             url_links.append(s)
    #         else:
    #             topics.append(s)

    chat = ChatOpenAI(model=GPT_MODEL)
    messages = [
        SystemMessage(
            content="You are an expert AI, tasked with generating accurate and complete frameworks to aid in human decision-making.  You only return python lists."
        ),
        HumanMessage(
            # content=f"Please make a bulleted list of the 4-8 most important criteria for {question}. {framework_context}. Do not number the list, or put parenthesis around any scores. Please also add to each criteria a number between 1 and 10 indicating how important this criteria is, where 1 is not important and 10 is very important. Please include a short explanation of each criteria.  Please return the criteria as a python list, with each item being a dictionary with keys for the name, weight, and description. The keys should be lowercase, and the dictionary should be valid python. Do not include any other commentary or code fences.  Explanations should be concise, no more than one sentence, and phrased as a question where a yes would score highly and a no would score lowly."
            content=f"""Please make a list of the 4-8 most important criteria for {question}. {framework_context}. Return a python list, where each item is a dictionary with keys for name, weight, and description.

Name should be a short summary of the criteria, between 1-4 words, and title-cased.
Explanations should be concise, no more than one sentence, and phrased as a question where a yes would score highly and a no would score lowly.
Weight should be a number between 1 and 10 indicating how important this criteria is, where 1 is not important and 10 is very important.

Please return the criteria as a python list, with each item being a dictionary with keys for the name, weight, and description. The keys should be lowercase, and the dictionary should be valid python. Ensure that any quotes inside dictionary values are properly escaped. Do not include any other commentary or code fences. Do not include examples."""
        ),
    ]

    chat_response = chat.invoke(messages).content

    # print(chat_response)

    # cool_chat = ChatOpenAI(temperature=0.3, model=GPT_MODEL)
    messages = [
        SystemMessage(
            content="You are an expert AI, tasked with generating accurate and complete frameworks to aid in human decision-making."
        ),
        HumanMessage(
            content=f"Please return a title for the following framework, used to analyze {question}.  Please only return the title, without any accompanying text.  Please keep it short - no more than four words.\n {chat_response}"
        ),
    ]

    title = chat.invoke(messages).content
    if title[0] == '"':
        title = title[1:]
    if title[-1] == '"':
        title = title[:-1]
    if title[-1] == ".":
        title = title[:-1]

    # print(title)
    # raw_criteria = chat_response.split("\n")
    # for c in raw_criteria:
    #     if c and c.strip() != "|":
    #         criteria.append(
    #             c.strip().replace("•", "").replace("•", "").replace("*", "")
    #         )
    # print(chat_response)
    try:
        criteria = ast.literal_eval(
            chat_response[chat_response.find("[") : chat_response.rfind("]") + 1]
        )
    except:
        print(chat_response)
        traceback.print_exc()

    topic_criteria_pairs = []
    criteria_with_weights = []
    # weighted_score = 0
    # max_weighted_score = 0
    topic = "temp"
    index = 0
    for c in criteria:
        if c:
            topic_criteria_pairs.append([topic, c, question, framework_context, title])

            # name = (
            #     c.split(":")[0]
            #     .strip()
            #     .replace("•", "")
            #     .replace("•", "")
            #     .replace("*", "")
            # )
            # weight = int(c.split("|")[0].split(":")[1].strip())
            # description = c.split("|")[1].strip()
            name = c["name"]
            weight = c["weight"]
            description = c["description"]
            criteria_with_weights.append(
                {
                    "name": name,
                    "weight": weight,
                    "description": description,
                    "index": index,
                }
            )
            index += 1
    # print(criteria_with_weights)
    data = {
        "framework": {
            "name": title,
            "criteria": criteria_with_weights,
            "description": f"OxGPT's first pass at a framework to help you {question}.",
            # "description": f"OxGPT's first pass at a framework to help you Add, remove, and change criteria, the title, and this description as you see fit.",
        },  # noqa
        "event_type": "generate-framework",
    }
    return_oxgpt(data, user_channel_id, gpt_guid)


@current_app.task
def oxgpt_analyze_subjects(data, user_channel_id):
    # print("oxgpt_analyze_subjects")
    criteria = []
    criteria = data["criteria"]
    # print(criteria)
    subjects = data["subjects"].split("\n")
    question = data["topic"]
    source_general_knowledge = data["sourceGeneralKnowledge"]
    source_search = data["sourceSearch"]
    source_news = data["sourceNews"]
    source_custom = data["sourceCustom"]
    source_text = data["sourceText"]
    gpt_guid = data["gptInstanceGUID"]
    files = data["fileContext"]
    source_urls = data["sourceUrls"]
    scoring_context = data["scoring_context"]
    # framework_context = data["topic_context"]
    full_context = {
        "source_general_knowledge": source_general_knowledge,
        "source_search": source_search,
        "source_news": source_news,
        "source_custom": source_custom,
        "source_urls": source_urls,
        "scoring_context": scoring_context,
    }
    subject_list = [s.strip() for s in subjects]
    url_links = []
    url_contents = []
    url_titles = []
    topics = []
    scorecards = {}

    for s in subject_list:
        if s.strip():
            if "://" in s:
                url_links.append(s)
            else:
                topics.append(s)

    with ThreadPoolExecutor(max_workers=20) as executor:
        results = list(executor.map(fetch_page_text, url_links))

    for r in results:
        if r:
            url, content, title = r
            url_contents[url] = content
            url_titles[url] = title

    for u in url_links:
        url = u.strip()
        # print(url)
        if url:
            # for unwanted in soup.find_all(['script', 'style', 'img', 'video', 'audio', 'aside', 'figure', 'footer', 'header', 'nav']):
            #     unwanted.extract()  # remove these tags and their contents
            #     url_contents[url] = soup.get_text()
            # /opt/homebrew/bin/wget -q -E -H -k -K -p -N --exclude-domains googletagmanager.com,cloudflareinsights.com,platform.linkedin.com --adjust-extension -e robots=off $URL
            # /opt/homebrew/bin/links -g -width 120 -codepage utf-8 -force-html -dump "index.html"
            # command = f"elinks -dump-charset utf-8 -force-html -dump {url}"
            command = f"links -g -width 120 -codepage utf-8 -force-html -dump {url}"
            process = subprocess.Popen(command, stdout=subprocess.PIPE, shell=True)
            output, error = process.communicate()

            if error:
                print(f"An error occurred: {error}")
            else:
                url_content = output.decode()  # decode bytes to string
                url_contents[url] = url_content

            # Get the title
            # command = f" curl -s {url} | grep -o '<title>[^<]*</title>' | sed 's/<title>\\(.*\\)<\\/title>/\1/g'"
            # process = subprocess.Popen(command, stdout=subprocess.PIPE, shell=True)
            # output, error = process.communicate()
            resp = requests.get(url)
            soup = BeautifulSoup(resp.content, "html.parser")
            title = soup.title.string if soup.title else url

            # title = output.decode()
            url_titles[url] = title

    # raw_criteria = chat_response.split("\n")
    # for c in raw_criteria:
    #     if c and c.strip() != "|":
    #         criteria.append(c.strip().replace("•", "").replace("•", "").replace("*", ""))

    # print(criteria)
    # Per topic method
    # topic_criteria_pairs = []
    # for topic in topics:
    #     scores = "".join([scores, f"{topic}\n"])
    #     topic = topic.strip()
    #     title = topic
    #     scorecards[topic] = {
    #         "name": title,
    #         "scores": [],
    #         "topic": topic,
    #         "ox_score": 0,
    #     }
    #     weighted_score = 0
    #     max_weighted_score = 0
    #     for c in criteria:
    #         if c:
    #             topic_criteria_pairs.append(
    #                 [topic, c, question, scoring_context, title]
    #             )

    #             name = c["name"]
    #             weight = c["weight"]

    # All topics scored together method
    # print(subjects)

    subjects_string = ", ".join(subjects)
    # print(subjects_string)
    context_messages = []
    urls_used = []

    use_precise = False
    if source_search or source_news:
        chat = ChatOpenAI(model=GPT_MODEL)
        messages = [
            HumanMessage(
                # content=f"Please make a bulleted list of the 4-8 most important criteria for {question}. {framework_context}. Do not number the list, or put parenthesis around any scores. Please also add to each criteria a number between 1 and 10 indicating how important this criteria is, where 1 is not important and 10 is very important. Please include a short explanation of each criteria.  Please return the criteria as a python list, with each item being a dictionary with keys for the name, weight, and description. The keys should be lowercase, and the dictionary should be valid python. Do not include any other commentary or code fences.  Explanations should be concise, no more than one sentence, and phrased as a question where a yes would score highly and a no would score lowly."
                content=f"""Is "{question} for {subjects_string}" a specific item, object, or idea that would be better to search for using precise, a quote-based query, or a general idea or concept that would be better as a semantic search?  Please return either PRECISE or SEMANTIC only, with no other text."""
            ),
        ]
        chat_response = chat.invoke(messages).content
        if "precise" in chat_response.lower():
            use_precise = True
    # print(f"using precise path? {use_precise}")
    if source_general_knowledge:
        pass
        # TODO
    if source_text:
        content = f"Here is some important context for the question that will be asked below.  It comes directly from me, and it should be treated as accurate and reliable.  Please prioritize this information when answering the question.:\n{source_text}"
        context_messages.append(content)
        data = {
            "message": "Evaluating custom text...",
            "complete": True,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)

    if files:
        for name, file in files.items():
            content = f"Here is some important context for the question that will be asked below.  It comes directly from me and files I added, and it should be treated as accurate and reliable.  Please prioritize this information when answering the question.:\n{file['text']}"
            context_messages.append(content)
            data = {
                "message": f"Evaluating {file['name']}",
                "complete": True,
                "event_type": "analyze-status-update",
            }
            return_oxgpt.delay(data, user_channel_id, gpt_guid)

    # print(source_custom)
    # print(source_urls)
    if source_custom and source_urls:
        data = {
            "message": "Evaluating custom urls...",
            "complete": False,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)
        cleaned_url_list = []
        for u in source_urls.split("\n"):
            # u = u.replace("https://", "")
            # if "/" in u:
            #     cleaned_url_list.append(u[:u.find("/")])
            # else:
            if u:
                cleaned_url_list.append(u)
        for url in cleaned_url_list:
            # print(url)
            # Find similar
            try:
                # similar_response = metaphor.find_similar(url, num_results=5)
                page_response = metaphor.search(
                    url, type="keyword", num_results=1
                ).get_contents()

                # contents_response = search_response.get_contents()
                # Print content for each result
                for content in page_response.contents:
                    # print(content.url)
                    for c in metaphor.get_contents(content.id).contents:
                        data = {
                            "message": f"Evaluating <a href='{u}' target='_blank'>{u}...</a>",
                            "complete": False,
                            "inset": True,
                            "event_type": "analyze-status-update",
                        }
                        return_oxgpt.delay(data, user_channel_id, gpt_guid)
                        page_contents = strip_tags(c.extract)
                        # print(page_contents)
                        # print(f"Title: {content.title}\nURL: {content.url}ID: {content.id}\nContent:\n{page_contents}\n")
                        content = f"Here is some accurate, important context for the question that will be asked below.  It comes from a reliable webpage that I specifically asked to include at {c.url}, and should be prioritized above more general information in the analysis. Please use it when responding to future questions, prompts, and analysis: \n\nTitle: {c.title}:\n{page_contents}"
                        context_messages.append(content)
                        urls_used.append(c.url)
                        u = c.url
                        if "?" in u:
                            u = u[: u.find("?")]
                        data = {
                            "message": f"Evaluating <a href='{u}' target='_blank'>{u}...</a>",
                            "complete": True,
                            "inset": True,
                            "event_type": "analyze-status-update",
                        }
                        return_oxgpt.delay(data, user_channel_id, gpt_guid)

            except:
                traceback.print_exc()
                print(url)
                print(page_response)
                print(page_response.__dict__)
        data = {
            "message": "Evaluating custom urls...",
            "complete": True,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)

    if source_news:
        # Do search on metaphor
        # print(
        #     f"Here is a recent news article about {question}, focused on these subjects: {subjects_string}: "
        # )
        search_response = None
        data = {
            "message": "Finding news...",
            "complete": False,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)
        if use_precise:
            try:
                search_response = metaphor.search(
                    # f"Here is a recent news article about '{question}' about {subjects_string}: ",
                    f"recent news about {question} for {subjects_string}",
                    type="keyword",
                    num_results=5,
                    start_published_date=(
                        datetime.datetime.now() - datetime.timedelta(days=60)
                    ).strftime("%Y-%m-%d"),
                )
            except:
                traceback.print_exc()
                print(
                    f"Error doing precise search for news about {question} for {subjects_string}"
                )

        else:
            try:
                full_domains = NEWS_DOMAINS
                news_chat = ChatOpenAI(model=GPT_MODEL)
                messages = [
                    HumanMessage(
                        # content=f"Please make a bulleted list of the 4-8 most important criteria for {question}. {framework_context}. Do not number the list, or put parenthesis around any scores. Please also add to each criteria a number between 1 and 10 indicating how important this criteria is, where 1 is not important and 10 is very important. Please include a short explanation of each criteria.  Please return the criteria as a python list, with each item being a dictionary with keys for the name, weight, and description. The keys should be lowercase, and the dictionary should be valid python. Do not include any other commentary or code fences.  Explanations should be concise, no more than one sentence, and phrased as a question where a yes would score highly and a no would score lowly."
                        content=f"""Please return ten URLs for the most reliable news websites you know about for information about "{question}" ({scoring_context}).  Don't include any descriptions or other information, or disclaimers about the knowledge cutoff.  Only reply with URLs, separated by line breaks."""
                    ),
                ]

                news_domains = news_chat.invoke(messages).content
                # print(news_domains)
                for n in news_domains.split("\n"):
                    if n.strip() and "." in n and "://" in n:
                        url = n.strip()[n.find("//") + 2 :]
                        if url.find("/") != -1:
                            url = url[: url.find("/")]
                        full_domains.append(url)

                search_response = metaphor.search(
                    # f"Here is a recent news article about '{question}' about {subjects_string}: ",
                    f"news about {question} and {subjects_string}",
                    num_results=5,
                    include_domains=full_domains,
                    start_published_date=(
                        datetime.datetime.now() - datetime.timedelta(days=60)
                    ).strftime("%Y-%m-%d"),
                )
            except:
                traceback.print_exc()
                print(
                    f"Error doing non-precise search for news about {question} for {subjects_string}"
                )
                print(full_domains)

        for result in search_response.results:
            u = result.url
            if "?" in u:
                u = u[: u.find("?")]

        if search_response:
            try:
                contents_response = search_response.get_contents()
            except:
                print("Error getting search_response.get_contents(), retrying...")
                contents_response = search_response.get_contents()

            # print(contents_response)
            # Print content for each result

            for content in contents_response.contents:
                data = {
                    "message": f"Evaluating <a href='{content.url}' target='_blank'>{content.url}...</a>",
                    "complete": False,
                    "inset": True,
                    "event_type": "analyze-status-update",
                }
                return_oxgpt.delay(data, user_channel_id, gpt_guid)
                # stripped_content = strip_tags(content.extract)
                # c = metaphor.get_contents(content.id)
                # print(content)
                # print(content.extract)
                stripped_content = strip_tags(content.extract)[:PAGE_CONTENT_CUTOFF]
                # print(stripped_content)
                # stripped_content = strip_tags(metaphor.get_contents(content.id).extract)
                # print(f"Title: {content.title}\nURL: {content.url}\nId: {content.id}\nContent:\n{stripped_content}\n")
                score_content = f"Here is accurate, true context for the question that will be asked below. It comes from a recent news article at {content.url}.  Please use it when responding to future questions, prompts, and analysis. \n\nTitle: {content.title}\n{stripped_content}"
                context_messages.append(score_content)
                urls_used.append(content.url)
                data = {
                    "message": f"Evaluating <a href='{content.url}' target='_blank'>{content.url}...</a>",
                    "complete": True,
                    "inset": True,
                    "event_type": "analyze-status-update",
                }
                return_oxgpt.delay(data, user_channel_id, gpt_guid)
        data = {
            "message": "Finding news...",
            "complete": True,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)

    if source_search:
        # Do search on metaphor
        # url_content = "test"
        data = {
            "message": "Finding good online sources...",
            "complete": False,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)
        search_response = None
        if use_precise:
            # print(f"{question} for {subjects_string}")
            try:
                search_response = metaphor.search(
                    f"{question} for {subjects_string}",
                    type="keyword",
                    num_results=3,
                )

            except:
                traceback.print_exc()
                print(
                    f"Error doing non-precise keyword search for {question} for {subjects_string}"
                )
        else:
            chat = ChatOpenAI(model=GPT_MODEL)
            messages = [
                HumanMessage(
                    # content=f"Please make a bulleted list of the 4-8 most important criteria for {question}. {framework_context}. Do not number the list, or put parenthesis around any scores. Please also add to each criteria a number between 1 and 10 indicating how important this criteria is, where 1 is not important and 10 is very important. Please include a short explanation of each criteria.  Please return the criteria as a python list, with each item being a dictionary with keys for the name, weight, and description. The keys should be lowercase, and the dictionary should be valid python. Do not include any other commentary or code fences.  Explanations should be concise, no more than one sentence, and phrased as a question where a yes would score highly and a no would score lowly."
                    content=f"""Please return ten URLs for the most reliable websites you know about for information about "{question}"?  Don't include any descriptions or other information, or disclaimers about the knowledge cutoff.  Only reply with URLs, separated by line breaks."""
                ),
            ]
            chat_response = chat.invoke(messages).content
            url_list = chat_response[chat_response.find("http") :].split("\n")
            cleaned_url_list = []
            for u in url_list:
                u = u.replace("https://", "")
                if "/" in u:
                    cleaned_url_list.append(u[: u.find("/")])
                else:
                    cleaned_url_list.append(u)
            # print(cleaned_url_list)
            # f"Here is a recent, accurate article about {question} about {subjects_string}:",
            # print(f"{question} for {subjects_string}")
            try:
                search_response = metaphor.search(
                    f"{subjects_string}",
                    # type="neural",
                    num_results=7,
                    include_domains=cleaned_url_list,
                )
            except:
                traceback.print_exc()
                print(f"Error doing non-precise neural search for {subjects_string}")
                print(f"Domains: {cleaned_url_list}")

        if search_response:
            try:
                contents_response = search_response.get_contents()
            except:
                print("Error getting search_response.get_contents(), retrying...")
                contents_response = search_response.get_contents()
            # print(contents_response)
            # Print content for each result
            for content in contents_response.contents:
                # print(content.url)
                u = content.url
                if "?" in u:
                    u = u[: u.find("?")]
                data = {
                    "message": f"Evaluating <a href='{u}' target='_blank'>{u}...</a>",
                    "complete": False,
                    "inset": True,
                    "event_type": "analyze-status-update",
                }
                return_oxgpt.delay(data, user_channel_id, gpt_guid)
                # stripped_content = strip_tags(content.extract)
                # c = metaphor.get_contents(content.id)
                # print(content)
                # print(content.extract)
                stripped_content = strip_tags(content.extract)[:PAGE_CONTENT_CUTOFF]
                # print(stripped_content)
                # stripped_content = strip_tags(metaphor.get_contents(content.id).extract)
                # print(f"Title: {content.title}\nURL: {content.url}\nId: {content.id}\nContent:\n{stripped_content}\n")
                score_content = f"Here is accurate, true context for the question that will be asked below. It comes from a reliable online source at {content.url}.  Please use it when responding to future questions, prompts, and analysis. \n\nTitle: {content.title}\n{stripped_content}"
                context_messages.append(score_content)
                urls_used.append(content.url)
                data = {
                    "message": f"Evaluating <a href='{u}' target='_blank'>{u}...</a>",
                    "complete": True,
                    "inset": True,
                    "event_type": "analyze-status-update",
                }
                return_oxgpt.delay(data, user_channel_id, gpt_guid)
                # print(score_content)
        data = {
            "message": "Finding good online sources...",
            "complete": True,
            "event_type": "analyze-status-update",
        }
        return_oxgpt.delay(data, user_channel_id, gpt_guid)

        # Per-subject information.
        for s in subjects:
            data = {
                "message": f"Researching {s}...",
                "complete": False,
                "event_type": "analyze-status-update",
            }
            return_oxgpt.delay(data, user_channel_id, gpt_guid)
            try:
                search_response = metaphor.search(
                    f"up-to-date, machine-readable information about {s}",
                    type="keyword",
                    num_results=3,
                )
                try:
                    contents_response = search_response.get_contents()
                except:
                    contents_response = search_response.get_contents()
                for content in contents_response.contents:
                    try:
                        u = content.url
                        if "?" in u:
                            u = u[: u.find("?")]
                        data = {
                            "message": f"Evaluating <a href='{u}' target='_blank'>{u}...</a>",
                            "complete": False,
                            "inset": True,
                            "event_type": "analyze-status-update",
                        }
                        return_oxgpt.delay(data, user_channel_id, gpt_guid)
                        stripped_content = strip_tags(content.extract)[
                            :PAGE_CONTENT_CUTOFF
                        ]
                        score_content = f"Here is up-to-date, accurate information about {s}. It comes from a reliable online source at {content.url}.  Please use it when responding to future questions, prompts, and analysis. \n\nTitle: {content.title}\n{stripped_content}"
                        context_messages.append(score_content)
                        urls_used.append(content.url)
                        data = {
                            "message": f"Evaluating <a href='{u}' target='_blank'>{u}...</a>",
                            "complete": True,
                            "inset": True,
                            "event_type": "analyze-status-update",
                        }
                        return_oxgpt.delay(data, user_channel_id, gpt_guid)
                        # print(content.url)
                    except:
                        print(
                            f"Error getting subject information about {s} from {content.url}"
                        )
                        traceback.print_exc()
            except:
                print(f"Error searching for {s}")
                traceback.print_exc()
            data = {
                "message": f"Researching {s}...",
                "complete": True,
                "event_type": "analyze-status-update",
            }
            return_oxgpt.delay(data, user_channel_id, gpt_guid)

        # Path for pure search, not super reliable.
        # print(f"Find information about these subjects: {subjects_string}, in the context of {question}")
        # search_response = metaphor.search(f"Find information about these subjects: {subjects_string}, in the context of {question}", num_results=5)

        # contents_response = search_response.get_contents()
        # # Print content for each result
        # for content in contents_response.contents:
        #     stripped_content = strip_tags(content.extract)
        #     # stripped_content = strip_tags(metaphor.get_contents(content.id).extract)
        #     print(f"Title: {content.title}\nURL: {content.url}\nId: {content.id}\nContent:\n{stripped_content}")
        #     content = f"Here is some context for the question that will be asked below.  Please use it when responding to future questions, prompts, and analysis: \n\nTitle: {content.title}\nURL: {content.url}\nContent:\n{stripped_content}"
        #     context_messages.append(content)
    data = {
        "message": "Compiling and evaluating research...",
        "complete": False,
        "event_type": "analyze-status-update",
    }
    return_oxgpt.delay(data, user_channel_id, gpt_guid)

    context_summary_response = gpt_summarize(
        context_messages,
        question=question,
        user_channel_id=user_channel_id,
        gpt_guid=gpt_guid,
    )

    data = {
        "message": "Compiling and evaluating research...",
        "complete": True,
        "event_type": "analyze-status-update",
    }
    return_oxgpt.delay(data, user_channel_id, gpt_guid)

    full_context["context_messages"] = [
        context_summary_response,
    ]
    full_context["user_channel_id"] = user_channel_id
    full_context["gpt_guid"] = gpt_guid

    full_context["urls_used"] = []
    for u in list(set(urls_used)):
        if "?" in u:
            u = u[: u.find("?")]
        u = u.strip()
        if u not in full_context["urls_used"]:
            full_context["urls_used"].append(u)

    # print(full_context["urls_used"])

    topic_criteria_pairs = []
    for c in criteria:
        if c:
            # for t in topics:
            #     topic_criteria_pairs.append([[t,], c, question, full_context])
            topic_criteria_pairs.append([topics, c, question, full_context])
    # print("topic_criteria_pairs")
    # print(topic_criteria_pairs)
    # score_criteria = partial(score_single_criteria, url=url, url_content=url_content)
    data = {
        "message": "Scoring...",
        "complete": False,
        "event_type": "analyze-status-update",
    }
    return_oxgpt.delay(data, user_channel_id, gpt_guid)

    thread_scoring_response_pairs = []
    max_workers = 6
    try:
        # Catch hangs - we should typically have all responses back in like 5 seconds.
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            thread_scoring_response_pairs = list(
                executor.map(
                    score_single_criteria_all_subjects,
                    # score_single_criteria_single_subjects,
                    topic_criteria_pairs,
                    timeout=180,
                )
            )
    except:
        print("Handling timeout errors, trying to recover.")
        missing_pairs = []
        # print("thread_scoring_response_pairs")
        # print(thread_scoring_response_pairs)
        for tcp in topic_criteria_pairs:
            found = False
            for rp in thread_scoring_response_pairs:
                # print(f"rp: {rp}\n  tcp: {tcp}")
                if rp[0] == tcp[0]:
                    found = True
                    break
            if not found:
                missing_pairs.append(tcp)

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            missing_thread_scoring_response_pairs = list(
                executor.map(score_single_criteria_all_subjects, missing_pairs)
                # executor.map(score_single_criteria_single_subjects, missing_pairs)
            )
        for mp in missing_thread_scoring_response_pairs:
            thread_scoring_response_pairs.append(mp)
    data = {
        "message": "Scoring...",
        "complete": True,
        "event_type": "analyze-status-update",
    }
    return_oxgpt.delay(data, user_channel_id, gpt_guid)

    # print(thread_scoring_response_pairs)
    scorecards = {}
    modified_at_ms = datetime.datetime.now().timestamp() * 1000
    # print(modified_at_ms)
    try:
        user = User.objects.get(ox_id=user_channel_id)
        oxgpt_id = user.ox_id
        oxgpt_name = user.full_name
    except:
        oxgpt_id = user_channel_id
        oxgpt_name = "Anonymous User"

    # print(thread_scoring_response_pairs)
    ox_scores = {}
    for topic in topics:
        topic = topic.replace("\n", "").strip()
        weighted_score = 0
        max_weighted_score = 0
        for r in thread_scoring_response_pairs:
            if r:
                if topic not in scorecards:
                    scorecards[topic] = {
                        "name": topic,
                        "scores": [],
                        "topic": topic,
                        "ox_score": 0,
                    }
                crit, scoring_response = r
                # print(crit, scoring_response)
                weight = crit["weight"]
                for t, topic_score in scoring_response.items():
                    # print(f"t: {t}, topic: {topic}, score: {topic_score}")
                    if topic == t:
                        raw_score = topic_score.get("score", 0)
                        comment = topic_score.get("comment", "")
                        # raw_skipped = topic_score.get("skipped")
                        skipped = topic_score.get("skipped")
                        speculation = topic_score.get("speculation")
                        separate_speculation = topic_score.get("separate_speculation")

                        if (speculation or speculation == 0) and (
                            separate_speculation or separate_speculation == 0
                        ):
                            nonspeculative = 5 - (
                                (speculation + separate_speculation) / 2
                            )
                        else:
                            # print("Missing separate_speculation")
                            nonspeculative = 5

                        # Weight half the score based on how speculative it is.
                        # /8 instead of /10 to give slight weighting to the poles, since GPT generally avoids them.
                        score = round(
                            (raw_score * 0.75) + (raw_score / 2 * nonspeculative / 8)
                        )
                        # print(f"{t}: {crit['name']}:  score: {score}, speculation: {speculation}, separate_speculation: {separate_speculation}")
                        # print("scoring for skipped:")
                        # print((raw_score/2))
                        # print((raw_score/4))
                        # print((raw_score/2 * nonspeculative / 10))
                        # print(f"score: {score}")
                        if score > 10:
                            score = 10
                        if (
                            (score < 1 and nonspeculative < 3)
                            or (score < 2 and nonspeculative < -4)
                            or nonspeculative < -3
                        ):
                            skipped = True
                            # score = 0
                        # print(f"skipped: {skipped}")
                        scorecards[topic]["scores"].append(
                            {
                                "name": crit["name"],
                                "criteria": {
                                    "name": crit["name"],
                                    "description": crit["description"],
                                    "weight": weight,
                                },
                                "score": score,
                                "gpt_score": score,
                                "gpt_scored_last": True,
                                "skipped": skipped,
                                # "comment": f"OxGPT comment: {comment}\n(Raw score: {raw_score} / Independent speculation score: {separate_speculation} / Integrated speculation score: {speculation} / final score: {score} / raw skipped: {raw_skipped})",
                                "comment": f"OxGPT comment: {comment}",
                            }
                        )
                        weighted_score += weight * int(score)
                        max_weighted_score += weight * 10
                        # print(score, weight, weighted_score, max_weighted_score)
                        scorecards[topic]["scorer"] = {
                            "id": oxgpt_id,
                            "full_name": oxgpt_name,
                            "modified_at_ms": modified_at_ms,
                        }
                        scorecards[topic]["framework"] = {
                            "name": topic,
                            "criteria": criteria,
                            "id": 1,
                        }
                        scorecards[topic]["modified_at_ms"] = modified_at_ms
                        scorecards[topic]["created_at_ms"] = modified_at_ms

        if max_weighted_score > 0:
            ox_score = 100 * weighted_score / max_weighted_score
            ox_scores[topic] = ox_score
    for index, sc in scorecards.items():
        for topic, score in ox_scores.items():
            if topic in index:
                sc["ox_score"] = ox_score

    reviewed_scorecards = scorecards

    scorecard_list = []
    for name, sc in reviewed_scorecards.items():
        scorecard_list.append(sc)
    sorted_scorecards = sorted(
        scorecard_list, key=lambda x: x["ox_score"], reverse=True
    )
    # print(sorted_scorecards)
    # print(full_context["urls_used"])
    data = {
        "scorecards": sorted_scorecards,
        "urls_used": full_context["urls_used"],
        "event_type": "analyze-subjects",
    }
    # print("data")
    # print(data)
    return_oxgpt(data, user_channel_id, gpt_guid)


@current_app.task
def oxgpt_generate_subjects(data, user_channel_id):
    # print("oxgpt_analyze_subjects")

    smart_generation = False

    question = data["topic"]
    framework_context = data["topic_context"]
    gpt_guid = data["gptInstanceGUID"]
    # previous_criteria = data["criteria"]
    # print(previous_criteria)
    # criteria_list = ""
    # for c in previous_criteria:
    #     criteria_list += f"- {c['name']}\n"

    if smart_generation:
        chat = ChatOpenAI(model=GPT_LARGE_CONTEXT_MODEL)
        messages = [
            SystemMessage(
                content="You are a creative AI, tasked with generating options for a given question, by using the context below."
            ),
        ]
        summarize_chat = ChatOpenAI(model=GPT_LARGE_CONTEXT_MODEL)
        year = datetime.datetime.now().year
        search_string = summarize_chat(
            [
                HumanMessage(
                    content=f"""Here is a framework topic and context.  The year is {year}.  If I was searching for specific, up-to-date options online (not general ideas), in twelve or fewer words, what would I search for?
    Topic: {question}
    Context: {framework_context}
    """
                ),
            ]
        ).content

        #
        # print(search_string)
        search_response = metaphor.search(
            f"{search_string}",
            num_results=3,
            type="keyword",
            use_autoprompt=False,
            start_published_date=(
                datetime.datetime.now() - datetime.timedelta(days=360)
            ).strftime("%Y-%m-%d"),
        )

        try:
            contents_response = search_response.get_contents()
        except:
            contents_response = search_response.get_contents()
        # print(contents_response)
        # Print content for each result
        for content in contents_response.contents:
            # stripped_content = strip_tags(content.extract)
            # c = metaphor.get_contents(content.id)
            # print(content)
            # print(content.extract)
            stripped_content = strip_tags(content.extract)[:PAGE_CONTENT_CUTOFF]
            if stripped_content:
                # print(stripped_content)
                summarize_chat = ChatOpenAI(
                    temperature=0.2, model=GPT_LARGE_CONTEXT_MODEL
                )
                summarized_content = summarize_chat(
                    [
                        HumanMessage(
                            content=f"""Please generate a list of the most relevant specific subjects (not general ideas) relating to {question} mentioned in this article: {stripped_content}. List no more than six."""
                        ),
                    ]
                ).content
                # print("summarized_content")
                # print(summarized_content)
                # print("summarized_content above")
                messages.append(
                    HumanMessage(
                        content=f"Here are some more recent potential options:\n {summarized_content}"
                    )
                )

        search_response = metaphor.search(
            search_string,
            num_results=3,
            type="keyword",
            use_autoprompt=False,
        )

        try:
            contents_response = search_response.get_contents()
        except:
            contents_response = search_response.get_contents()
        # print(contents_response)
        # Print content for each result
        for content in contents_response.contents:
            # stripped_content = strip_tags(content.extract)[:PAGE_CONTENT_CUTOFF]
            # c = metaphor.get_contents(content.id)
            # print(content)
            # print(content.extract)
            stripped_content = strip_tags(content.extract)[:PAGE_CONTENT_CUTOFF]
            if stripped_content:
                # print(stripped_content)
                # summarize_chat = ChatOpenAI(temperature=0.2, model=GPT_LARGE_CONTEXT_MODEL)
                summarized_content = summarize_chat(
                    [
                        HumanMessage(
                            content=f"""Please generate a list of the most relevant specific subjects (not general ideas) relating to {question} mentioned in this article: {stripped_content}. List no more than six."""
                        ),
                    ]
                ).content
                # print("summarized_content")
                # print(summarized_content)
                # print("summarized_content above")
                messages.append(
                    HumanMessage(
                        content=f"Here are some potential options:\n {summarized_content}"
                    )
                )

        messages.append(
            HumanMessage(
                content=f"""Generate between 2 and 6 specific options for the question of {question}. {framework_context}. Where possible, use the options above. Subjects should be short, singular topics, not sentences or descriptions.  Return valid code in the form of a python list, where each item is a the name of the subject. Use double-quotes for the list, and escape any double-quotes in the list items. Do not number the list."""
            ),
        )
    else:
        chat = ChatOpenAI(model=GPT_MODEL)
        messages = [
            SystemMessage(
                content="You are a creative AI, tasked with generating options for a given question, by using the context below."
            ),
        ]
        messages.append(
            HumanMessage(
                content=f"""Generate between 2 and 6 specific options for the question of {question}. {framework_context}. Subjects should be short, singular topics, not sentences or descriptions.  Return valid code in the form of a python list, where each item is a the name of the subject. Use double-quotes for the list, and escape any double-quotes in the list items. Do not number the list."""
            ),
        )

    # Do not number the list, or put parenthesis around any scores. Please also add to each criteria a number between 0 and 10 indicating how important this criteria is, where 1 is not important and 10 is very important. Please include a short explanation of each criteria.  Please return the criteria as a python list, with each item being a dictionary with keys for the name, weight, and description. The keys should be lowercase, and the dictionary should be valid python. Do not include any other commentary or code fences.  Explanations should be concise, no more than one sentence, and phrased as a question where a yes would score highly and a no would score lowly."
    chat_response = chat.invoke(messages).content
    # print(chat_response)
    chat_response = chat_response[
        chat_response.find("[") : chat_response.rfind("]") + 1
    ]
    new_subjects = [
        "Error generating subjects.  Feel free to try again!  If you see this message again, it might not be a topic that Ox knows enough about.",
    ]
    try:
        new_subjects = "\n".join(ast.literal_eval(chat_response))
    except:
        traceback.print_exc()

    data = {
        "subjects": new_subjects,
        "event_type": "generate-subjects",
    }
    return_oxgpt(data, user_channel_id, gpt_guid)


@current_app.task
def oxgpt_generate_more_criteria(data, user_channel_id):
    # Return all objects created or modified, and success true or false.
    # print(data)
    # criteria = []
    question = data["topic"]
    framework_context = data["topic_context"]
    previous_criteria = data["criteria"]
    gpt_guid = data["gptInstanceGUID"]
    # print(previous_criteria)
    criteria_list = ""
    for c in previous_criteria:
        criteria_list += f"- {c['name']}\n"

    chat = ChatOpenAI(model=GPT_MODEL)
    messages = [
        SystemMessage(
            content="You are an expert AI, tasked with generating accurate and complete frameworks to aid in human decision-making. You only return python lists."
        ),
        HumanMessage(
            content=f"""Generate three (3) more important criteria for {question}. {framework_context}. Please do not include any of the criteria in the following list {criteria_list}.

Return a python list, where each item is a dictionary with keys for name, weight, and description.

Name should be a short summary of the criteria, between 1-4 words, and title-cased.
Explanations should be concise, no more than one sentence, and phrased as a question where a yes would score highly and a no would score lowly.
Weight should be a number between 1 and 10 indicating how important this criteria is, where 1 is not important and 10 is very important.

Please return the criteria as a python list, with each item being a dictionary with keys for the name, weight, and description. The keys should be lowercase, and the dictionary should be valid python. Ensure that any quotes inside dictionary values are properly escaped. Do not include any other commentary or code fences. Do not include examples or code evaulation."""
        ),
    ]

    # Do not number the list, or put parenthesis around any scores. Please also add to each criteria a number between 0 and 10 indicating how important this criteria is, where 1 is not important and 10 is very important. Please include a short explanation of each criteria.  Please return the criteria as a python list, with each item being a dictionary with keys for the name, weight, and description. The keys should be lowercase, and the dictionary should be valid python. Do not include any other commentary or code fences.  Explanations should be concise, no more than one sentence, and phrased as a question where a yes would score highly and a no would score lowly."
    chat_response = chat.invoke(messages).content
    chat_response = chat_response[
        chat_response.find("[") : chat_response.rfind("]") + 1
    ]
    # print(chat_response)

    try:
        new_criteria = ast.literal_eval(chat_response)
    except:
        traceback.print_exc()
    # print("new_criteria")
    # print(new_criteria)
    # for c in new_criteria:
    #     if c and c.strip() != "|":
    #         criteria.append(
    #             c.strip()
    #             .replace("•", "")
    #             .replace("•", "")
    #             .replace("*", "")
    #             .replace("- ", "")
    #         )

    # topic_criteria_pairs = []
    # weighted_score = 0
    # max_weighted_score = 0
    # topic = "temp"
    criteria_with_weights = []
    index = previous_criteria[-1]["index"] + 1
    for c in new_criteria:
        if c:
            criteria_with_weights.append(
                {
                    "name": c["name"],
                    "weight": c["weight"],
                    "description": c["description"],
                    "index": index,
                }
            )
            index += 1
    # print(criteria_with_weights)
    data = {
        "criteria": criteria_with_weights,
        "event_type": "generate-more-criteria",
    }
    return_oxgpt(data, user_channel_id, gpt_guid)


@current_app.task
def export_pdf(data, user_channel_id):
    # print(f"oxgpt_generate_framework: {user_channel_id}")

    # data = {
    #     "report_id": report_id,
    #     "title": title,
    #     "org_name": org_name,
    #     "distribution_text": distribution_text,
    #     "page_theme": page_theme,
    #     "org_logo": org_logo,
    # }

    u = User.objects.get(ox_id=user_channel_id)
    blob = b""
    name = "Unknown"
    if "report_id" in data:
        obj = Report.authorized_objects.authorize(user=u).findable.get(
            ox_id=data["report_id"]
        )
        name = obj.name

        blob = obj.generate_pdf(
            title=data["title"],
            org_name=data["org_name"],
            distribution_text=data["distribution_text"],
            page_theme=data["page_theme"],
            org_logo=data["org_logo"],
            requesting_user=u,
            page_domain=data["page_domain"]
            # request=request,
        )
    elif "stack_id" in data:
        obj = Stack.authorized_objects.authorize(user=u).findable.get(
            ox_id=data["stack_id"]
        )
        name = obj.name

        blob = obj.generate_pdf(
            title=data["title"],
            org_name=data["org_name"],
            distribution_text=data["distribution_text"],
            page_theme=data["page_theme"],
            org_logo=data["org_logo"],
            requesting_user=u,
            page_domain=data["page_domain"]
            # request=request,
        )

    # resp["Content-Disposition"] = f"filename={r.name}.pdf"
    data = {
        "event_type": "export-pdf",
        "blob": base64.b64encode(blob).decode("utf-8"),
        "filename": f"{name}.pdf",
    }
    # print("exported pdf.")
    return return_export(data, user_channel_id)


@current_app.task
def oxgpt_analyze_file(data, user_channel_id):
    name = data["name"]
    size = data["size"]
    gpt_guid = data["gptInstanceGUID"]

    b64content = base64.b64decode(data["content"])
    if name.lower().endswith(".pdf"):
        file_type = "pdf"
        # PyMuPDF
        # pdf_buffer = fitz.open("pdf", b64content)

        # mu_text = ""
        # print(len(pdf_buffer))
        # for page_num in range(len(pdf_buffer)):
        #     page = pdf_buffer.load_page(page_num)
        #     mu_text += page.get_text("text")

        # PyPDF2
        pdf_buffer = io.BytesIO(b64content)
        pdf_reader = pypdf.PdfReader(pdf_buffer)
        pypdf_text = ""

        for page in pdf_reader.pages:
            pypdf_text += page.extract_text()

        # pdfminer
        # if len(text) < 10:
        #     print(text)
        #     print("falling back to pdfminer")
        # pdf_buffer.seek(0)
        # miner_text = extract_text(pdf_buffer)

        # pdfplumber
        # pdf_buffer.seek(0)
        # plumber_text = ""
        # with pdfplumber.open(pdf_buffer) as pdf:
        #     for page in pdf.pages:
        #         print(page.extract_text())
        #         plumber_text += page.extract_text()

        pdf_buffer.seek(0)
        images = convert_from_bytes(pdf_buffer.read())

        text = pypdf_text
        if len(text.split(" ")) < 10:
            tesseract_text = ""
            for i, image in enumerate(images):
                tesseract_text += image_to_string(image)

            if len(tesseract_text.split(" ")) > len(text.split(" ")):
                text = tesseract_text
            # print(f"tesseract_text: {len(tesseract_text.split(' '))}")

        # print(f"mu_text: {len(mu_text.split(' '))}")
        # print(f"pypdf_text: {len(pypdf_text.split(' '))}")
        # print(f"miner_text: {len(miner_text.split(' '))}")
        # print(f"plumber_text: {len(plumber_text.split(' '))}")

    elif name.lower().endswith(".doc"):
        file_type = "doc"
        # print("handle doc")
        docx_buffer = io.BytesIO(b64content)
        doc = Document(docx_buffer)

        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif name.lower().endswith(".docx"):
        file_type = "doc"
        # print("handle docx")
        docx_buffer = io.BytesIO(b64content)
        with docx2python(docx_buffer) as docx_content:
            text = docx_content.text

        # doc = Document(docx_buffer)

        # text = ""
        # for para in doc.paragraphs:
        #     text += para.text + "\n"
        # print(text)
    else:
        file_type = "other"
        text = b64content.decode("utf-8")

    num_words = len(text.split(" "))
    # print(name)
    # print(text)
    data = {
        "text": text,
        "filename": name,
        "type": file_type,
        "size": size,
        "num_words": num_words,
        "event_type": "analyze-file-contents",
    }
    return_oxgpt(data, user_channel_id, gpt_guid)
